"""Build cover letter for ESR submission + regenerated manuscript DOCX.

Usage:  python build_cover_letter.py
Outputs into ../manuscript/ :
  - cover_letter.docx
  - manuscript_en.docx  (overwritten, now reflects full revised content)
  - manuscript_ja.docx  (overwritten)
"""
from __future__ import annotations

import os
import re
import subprocess
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
MS = os.path.join(ROOT, "manuscript")

os.makedirs(MS, exist_ok=True)

# ---------------------------------------------------------------------------
# Font helpers
# ---------------------------------------------------------------------------

def set_run_font(run, name="Times New Roman", size=11, bold=False, italic=False,
                 color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), "MS Mincho")


def add_para(doc, text, size=11, bold=False, italic=False, alignment=None,
             space_after=6, space_before=0, first_line_indent=None,
             color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(first_line_indent)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading_styled(doc, text, level=1):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(text)
    set_run_font(run, size={1: 14, 2: 12, 3: 11}[level], bold=True)
    return h


# ---------------------------------------------------------------------------
# Cover letter
# ---------------------------------------------------------------------------

def build_cover_letter():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Date
    add_para(doc, "May 17, 2026", size=11)

    # Addressee
    add_para(doc, "Professor Yasuhide Okuyama and Professor Rosa Duarte", size=11)
    add_para(doc, "Editors-in-Chief", size=11)
    add_para(doc, "Economic Systems Research", size=11)

    add_para(doc, "", size=11)  # blank line

    # Subject
    add_para(doc, "Re: Submission of \"Time-Varying Time-to-Build in Capital Accounting: "
             "Reconciling Flow- and Stock-Based National Wealth Measures through a "
             "Time-Varying Perpetual Inventory Method with Intangible Capital\"",
             bold=True, size=11)

    add_para(doc, "", size=11)

    # Body
    body_paras = [
        "Dear Editors,",

        "I am pleased to submit the enclosed manuscript for consideration for "
        "publication in Economic Systems Research. The paper develops a framework "
        "for reconciling flow-based (GDP) and stock-based (national wealth) accounts "
        "by relaxing two standard assumptions in the perpetual inventory method: "
        "zero time-to-build and the exclusion of intangible capital.",

        "The paper makes four specific contributions that align with the journal's "
        "scope. First, it introduces a time-varying mean investment-to-output lag in "
        "a production-function PIM framework and shows that this improves out-of-sample "
        "GDP level predictions by 12–13% across 39 OECD and middle-income economies. "
        "Second, it jointly identifies the time-to-build parameter μ and the intangible "
        "capital share β against both production data (Penn World Table) and wealth "
        "data (World Bank CWON), demonstrating that the flow and stock accounts agree "
        "to within 1–2% for most advanced economies when both parameters are estimated "
        "rather than imposed. Third, it introduces a **relational PIM (RPIM)** — "
        "inspired by the Brass (1971) relational model in demography — which formalises "
        "the PIM-CWON consistency check through two diagnostic parameters (ρ₁, ρ₂) "
        "and provides a δ-ρ₂ sensitivity analysis to assess whether depreciation "
        "mis-specification could explain the residual gap. Fourth, it provides bootstrap "
        "confidence intervals showing that neither μ nor β is well identified from "
        "production data alone; the wealth-side constraint from the national accounting "
        "identity dW/dt = S(Y) − δW is essential.",

        "The paper was previously submitted to the Review of Income and Wealth, where "
        "it received a desk rejection. I have taken the Editor's detailed and "
        "constructive feedback seriously and made substantial revisions. The most "
        "important changes are: (1) the paper now engages explicitly with the asset-"
        "specific PIM literature (OECD Manual, vintage models, age-efficiency profiles); "
        "(2) the demographic analogy that motivated the framework is presented as a "
        "heuristic rather than a formal claim of equivalence, with a dedicated section "
        "discussing where the analogy holds and where it breaks down; (3) additional "
        "benchmark models (AR(1) distributed lag, broken-trend specification) have been "
        "added to test the robustness of the time-varying lag result; (4) the "
        "empirical results are reported more honestly, including the negative finding "
        "that intangible capital alone does not improve out-of-sample prediction; and "
        "(5) a new **relational PIM (M5)** has been introduced, inspired by the Brass "
        "(1971) relational model, which formalises the PIM-CWON consistency check through "
        "two diagnostic parameters (ρ₁, ρ₂) and provides a δ-ρ₂ sensitivity analysis "
        "to separate depreciation effects from time-to-build effects.",

        "I believe the revised manuscript is now a better fit for Economic Systems "
        "Research than for its previous target. The paper's emphasis on the accounting-"
        "identity linkage between production and wealth accounts, its cross-country "
        "empirical approach, and its practical recommendations for wealth accounting "
        "programmes (CWON, IWI) all speak directly to the journal's readership.",

        "This manuscript has not been published previously and is not under "
        "consideration elsewhere. I have no conflicts of interest to declare.",

        "Thank you for your time and consideration. I look forward to hearing from you.",
    ]

    for text in body_paras:
        add_para(doc, text, size=11)

    add_para(doc, "", size=11)

    # Signature
    add_para(doc, "Sincerely,", size=11)
    add_para(doc, "", size=6)
    add_para(doc, "Tatsuki Onishi", size=11)
    add_para(doc, "Data Science AI Innovation Research Promotion Center, Shiga University",
             size=10, italic=True)
    add_para(doc, "E-mail: bougtoir@gmail.com", size=10, italic=True)

    out = os.path.join(MS, "cover_letter.docx")
    doc.save(out)
    print("wrote", out)


# ---------------------------------------------------------------------------
# Rebuild manuscript docx from markdown with full content
# ---------------------------------------------------------------------------

def markdown_to_docx_simple(md_path, out_path, lang="en"):
    """Convert markdown to docx preserving headings, paragraphs, tables, code blocks."""
    with open(md_path, encoding="utf-8") as fh:
        text = fh.read()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    lines = text.split("\n")
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        # Code block handling
        if stripped.startswith("```"):
            if in_code_block:
                # End code block
                for cl in code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.left_indent = Cm(1)
                    run = p.add_run(cl if cl else " ")
                    set_run_font(run, name="Courier New", size=9)
                code_lines = []
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                code_lines = []
                i += 1
                continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Horizontal rule
        if stripped.startswith("---") and len(stripped) >= 3:
            i += 1
            continue

        # Headings
        if stripped.startswith("# ") or stripped.startswith("#　"):
            add_heading_styled(doc, stripped.lstrip("#　 ").strip(), 1)
        elif stripped.startswith("## ") or stripped.startswith("##　"):
            add_heading_styled(doc, stripped.lstrip("#　 ").strip(), 2)
        elif stripped.startswith("### ") or stripped.startswith("###　"):
            add_heading_styled(doc, stripped.lstrip("#　 ").strip(), 3)
        elif stripped.startswith("**Table"):
            # Bold table heading
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped)
            set_run_font(run, bold=True, size=10)
        elif stripped.startswith("| "):
            # Table row - handle simple markdown tables
            if "|" in stripped:
                cells = [c.strip() for c in stripped.split("|") if c.strip()]
                # Skip separator rows like |---|---|
                if all(re.match(r"^-+$", c) for c in cells):
                    i += 1
                    continue
                # We don't try to build actual docx tables from markdown table
                # rows; render as monospaced text
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(stripped)
                set_run_font(run, name="Courier New", size=9)
        elif stripped.startswith("["):
            # Placeholder like [Insert table 1 here] or [Figure 1 here]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            set_run_font(run, italic=True, size=10,
                         color=(128, 128, 128))
        elif stripped == "":
            # Blank line between paragraphs
            pass
        else:
            # Normal paragraph: clean markdown bold/italic markers
            text_clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            text_clean = re.sub(r"\*(.+?)\*", r"\1", text_clean)
            # Skip "**Keywords**: ..." etc that we already handled
            if text_clean.strip():
                add_para(doc, text_clean, size=11)

        i += 1

    doc.save(out_path)
    print("wrote", out_path)


def build_full_manuscripts():
    en_md = os.path.join(MS, "manuscript_en.md")
    ja_md = os.path.join(MS, "manuscript_ja.md")

    if os.path.exists(en_md):
        markdown_to_docx_simple(en_md, os.path.join(MS, "manuscript_en.docx"), lang="en")
    if os.path.exists(ja_md):
        markdown_to_docx_simple(ja_md, os.path.join(MS, "manuscript_ja.docx"), lang="ja")


def main():
    build_cover_letter()
    build_full_manuscripts()


if __name__ == "__main__":
    main()
